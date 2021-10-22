from docx import Document
import pdfplumber 
import fitz 
import io
from PIL import Image
import time

class Convertidor():
    
    def __init__(self,nombre_archivo_pdf,contine_indice=[]):
        self.nombre_archivo_pdf=nombre_archivo_pdf
        self.nombre_archivo_txt_provisorio=""
        self.archivo_txt_final=""
        self.imagenes=[]
        self.archivo_docx_final=""        
        self.contiene_indice=contine_indice
        self.indice_provisorio=[]
        self.indice=[]
        self.porcentajes_completados=[75,50,25]
        self.accion_funcion=""

    #decorador para calcular los tiempos de ejecucion de cada funcion
    def calcula_tiempos_arg(arg_dec=""):    
        def calcula_tiempos(funcion):
            def funcion_a_decorar(self,*args):  
                start=time.time()      
                funcion(self,*args)
                end=time.time() 
                tiempo=end-start   
                print(f"{arg_dec} en %.2f seconds." % tiempo) 
                if funcion.__name__== "txt_a_docx":
                    print(f"Finalizado el proceso para {self.nombre_archivo_pdf}")   
            return funcion_a_decorar    
        return calcula_tiempos

    #Imprime el porcentaje completado de la funcion que convierte pdf a txt provisorio
    def calcula_porcentaje_completado(self,total_paginas,pagina_actual):         
        try:
            if self.porcentajes_completados ==[]:
                pass
            else:
                if pagina_actual > self.porcentajes_completados[-1]:
                    print(f'Completado {self.porcentajes_completados[-1]}%')
                    self.porcentajes_completados.pop()
        except:
            pass   
            
    

    @calcula_tiempos_arg("Convirtiendo pdf a txt...")
    def pdf_a_txt(self):          
        f = open (self.nombre_archivo_pdf[:-4]+'_provisorio.txt','w')  
        #creo un archivo txt provisorio donde se iran guardando los textos y las posiciones de las imagenes      
        self.nombre_archivo_txt_provisorio=self.nombre_archivo_pdf[:-4]+'_provisorio.txt'  
        #libreria para obtener las imagenes de los archivos PDF
        with pdfplumber.open(self.nombre_archivo_pdf) as pdf:
            #libreria para obtener el texto de los archivos PDF
            with fitz.open(self.nombre_archivo_pdf) as my_pdf_file :
                for x,y in zip(pdf.pages,range(len(my_pdf_file)+1)): 
                    try:
                        if y+1 in self.contiene_indice:                           
                            texto=x.extract_text()  
                            self.indice_provisorio.append(texto)
                            continue
                        
                        self.calcula_porcentaje_completado(len(my_pdf_file),y+1)
                        texto=x.extract_text()  
                        page = my_pdf_file[y-1]
                        images = page.getImageList()  
                        if images :
                            f.write("\n #IMAGEN# \n") # si la pagina tiene imagenes, escribira en el txt #IMAGEN#, para tomar referencia donde se insertarán luego las mismas
                        
                        if texto != None:
                            f.write(texto+"\n")  
                        else:
                            f.write("\n")       
                        for image_number, image in enumerate(page.getImageList(), start=1):
                            xref_value = image[0]
                            base_image = my_pdf_file.extractImage(xref_value)
                            image_bytes = base_image["image"]
                            ext = base_image["ext"]
                            image = Image.open(io.BytesIO(image_bytes))
                            image.save(open(f"Page{y}Image{image_number}.{ext}", "wb"))   
                            self.imagenes.append(f"Page{y}Image{image_number}.{ext}")
                    except:
                        pass        
        f.close()       


    @calcula_tiempos_arg
    def calcula_indices(self):
        for x in range(len(self.indice_provisorio)):
            lista_separada=self.indice_provisorio[x].split("\n")
            self.indice+=lista_separada  


        
     
    @calcula_tiempos_arg("Ajustando txt...")
    def acomoda_txt(self):
        lineas_txt = []
        with open(self.nombre_archivo_txt_provisorio) as fname:
            lineas = fname.readlines()
            for linea in lineas:
                lineas_txt.append(linea.strip('\n'))  
        lista=""
        for x in enumerate(lineas_txt):            
            try:
                ultima_linea=lineas_txt[x[0]-1]
                penultima_linea=lineas_txt[x[0]-2]  
                #salteo las paginas informadas como indice              
                if self.indice and self.indice[-1] >= x[0]:
                    continue 
                #si la linea contiene ese caracter especial ,lo considero como posible titulo ,lo cambio a mayuscula y agrego espacios.
                if ".-" in x[1]:
                    cadena=lineas_txt[x[0]]   
                    cadena=cadena.upper()     
                    cadena_def="\n\n"+cadena+"\n\n"
                    lineas_txt[x[0]]=cadena_def
                    continue
                #si la linea entera esta en mayuscula ,la considero titulo 
                if x[1].isupper():
                    cadena=lineas_txt[x[0]]        
                    cadena_def="\n\n"+cadena+"\n\n"
                    lineas_txt[x[0]]=cadena_def
                    continue  
                #si la linea termina con punto ,lo considero como punto y a parte ,y le agrego salto de linea. 
                if x[1][-1]==".":
                    cadena=lineas_txt[x[0]]        
                    cadena_def=cadena+"\n"
                    lineas_txt[x[0]]=cadena_def   
                    continue
                #si la linea anterior esta en blanco y la actual no continue imagenes ,posiblemente sea un titulo
                if ultima_linea==""  and x[1]!= " #IMAGEN# ":                                        
                    cadena=lineas_txt[x[0]].upper() 
                    cadena_def="\n\n"+cadena+"\n\n"
                    lineas_txt[x[0]]=cadena_def  
                    continue 

                # si una linea completa es igual a una linea completa que esta marcada como indice, la tomo como titulo o capitulo.        
                if self.contiene_indice:                    
                    if x[0]>self.contiene_indice[-1]:
                        if x[1] in self.indice:                            
                            cadena=lineas_txt[x[0]].upper()                            
                            cadena_def="\n\n"+cadena+"\n\n"
                            lineas_txt[x[0]]=cadena_def  
                            continue             

            except:
                cadena=lineas_txt[x[0]]        
                cadena_def=cadena+"\n"
                lineas_txt[x[0]]=cadena_def 
                pass

        for dato in lineas_txt:              
            try: 
                if dato[-1]==" " and dato[-1].isupper()== False:
                    lista=lista + dato
                    continue
                if dato[-1]!=" " and dato[-1].isupper()== False:
                    lista=lista+" " + dato    
            except:
                pass
        #guardo txt modificado como final   
        f = open (self.nombre_archivo_pdf[:-4]+'_final.txt','w')        
        f.write(lista)
        f.close()       
        self.archivo_txt_final=    lista        



    @calcula_tiempos_arg('Conviertiendo txt a doc...')
    #genero el archivo docx a partir del txt final
    def txt_a_docx(self):        
        document = Document()
        txt_a_lista=self.archivo_txt_final.split("\n\n")
        for x in enumerate(txt_a_lista):
            try:                
                if x[1].isupper() and x[1]!=" #IMAGEN# " :                    
                    document.add_heading(x[1], level=1)                
                else:    
                    document.add_paragraph(x[1])
            except:
                document.add_paragraph(x[1])


        all_paras = document.paragraphs        
        contador=0        
        for i in all_paras:
            if i.text==" #IMAGEN# ":
                try:
                    i.text=""
                    r = i.add_run()
                    r.add_text=""
                    r.add_picture(self.imagenes[contador])        
                    contador+=1
                except:
                    pass

        self.archivo_docx_final=document        
        document.save(self.nombre_archivo_pdf[:-4]+'.docx')    



if __name__ == '__main__':
    conv=Convertidor("HACIA EL AUTOGOBIERNO - UNA CRITICA AL PODER POLITICO.pdf")
    conv.pdf_a_txt()
    conv.calcula_indices()    
    conv.acomoda_txt()
    conv.txt_a_docx()
